import re
from pathlib import Path

from docx import Document as DocxDocument

from app.rag.blocks import Block, ParsedDocument
from app.rag.parsers.base import Parser


class DocxParser(Parser):
    def parse(self, path: str, section_marker: re.Pattern[str] | None = None) -> ParsedDocument:
        # python-docx raises PackageNotFoundError for a missing file, which is
        # not a FileNotFoundError - the structure endpoint's "source file is no
        # longer available" 404 would degrade into a 500 without this.
        if not Path(path).is_file():
            raise FileNotFoundError(path)

        doc = DocxDocument(path)
        blocks: list[Block] = []
        current_section: str | None = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style is not None else ""
            if style.startswith("Heading") or style == "Title":
                current_section = text
                blocks.append(Block(text=text, block_type="heading", section=current_section))
            elif style.startswith("List"):
                blocks.append(Block(text=text, block_type="list_item", section=current_section))
            else:
                blocks.append(Block(text=text, block_type="paragraph", section=current_section))

        for table in doc.tables:
            # row.cells expands a merge to one entry per grid column (and per row
            # for a vertical merge), handing back the same underlying <w:tc>
            # repeatedly. Emitting each one indexes and retrieves the same text
            # several times - the same duplicate the HTML parser skips nested
            # tags to avoid. Per-table, not per-row, so vertical merges dedupe too.
            seen: set = set()
            for row in table.rows:
                for cell in row.cells:
                    if cell._tc in seen:
                        continue
                    seen.add(cell._tc)
                    text = cell.text.strip()
                    if text:
                        blocks.append(Block(text=text, block_type="table_cell", section=current_section))

        return ParsedDocument(blocks=blocks)
