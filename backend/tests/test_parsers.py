import pytest
from docx import Document as DocxDocument

from app.rag.parsers import get_parser
from app.rag.parsers.pdf_parser import _columns, _is_heading


def _write_pdf(path, pages_lines: list[list[str]]) -> None:
    """Minimal text-only PDF writer. No PDF-authoring library is installed and
    adding one just for fixtures is not worth it - the parsers must be proven
    against bytes pypdf actually reads, not against a mocked extract_text."""
    objs: dict[int, bytes] = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        3: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    }
    page_ids = [5 + 2 * i for i in range(len(pages_lines))]
    objs[2] = b"<< /Type /Pages /Count %d /Kids [%s] >>" % (
        len(pages_lines),
        b" ".join(b"%d 0 R" % p for p in page_ids),
    )
    for i, lines in enumerate(pages_lines):
        parts = [b"BT", b"/F1 12 Tf", b"1 0 0 1 72 720 Tm", b"16 TL"]
        for line in lines:
            escaped = line.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
            parts.append(b"(" + escaped.encode("latin-1") + b") Tj T*")
        stream = b"\n".join([*parts, b"ET"])
        objs[4 + 2 * i] = b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream)
        objs[5 + 2 * i] = (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 3 0 R >> >> /Contents %d 0 R >>" % (4 + 2 * i)
        )

    out = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += b"%d 0 obj\n" % num + objs[num] + b"\nendobj\n"
    xref_offset, size = len(out), max(objs) + 1
    out += b"xref\n0 %d\n0000000000 65535 f \n" % size
    for num in range(1, size):
        out += b"%010d 00000 n \n" % offsets.get(num, 0)
    out += b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n" % (size, xref_offset)
    path.write_bytes(bytes(out))


_UNESCAPED_CODES = [code for code in range(33, 127) if code not in (40, 41, 92)]


def _write_positioned_pdf(path, pages, page_height: int = 792) -> None:
    """A PDF whose text runs sit at absolute positions in a CHOSEN content-stream
    order, carrying arbitrary Unicode.

    `_write_pdf` above cannot express the failure this parser exists to fix: the
    numerals have to be drawn AFTER the Hangul run they belong inside. Hangul
    needs no embedded font here - pdfminer reads characters out of the
    /ToUnicode CMap, so Helvetica's codes can stand for any code point - and
    every glyph is given the same 500/1000 width, which makes the geometry the
    parser groups words and lines on exact rather than font-metric dependent.

    Each page is a list of (x, top, size, text) runs, drawn in the order given.
    """
    codes: dict[str, int] = {}
    for page in pages:
        for _, _, _, text in page:
            for character in text:
                if character not in codes:
                    codes[character] = _UNESCAPED_CODES[len(codes)]

    bfchar = b"".join(b"<%02X> <%04X>\n" % (code, ord(c)) for c, code in codes.items())
    cmap = (
        b"/CIDInit /ProcSet findresource begin\n12 dict begin\nbegincmap\n"
        b"/CMapName /Fixture def\n/CMapType 2 def\n"
        b"1 begincodespacerange\n<21> <7E>\nendcodespacerange\n"
        b"%d beginbfchar\n%sendbfchar\nendcmap\n"
        b"CMapName currentdict /CMap defineresource pop\nend\nend" % (len(codes), bfchar)
    )

    objs: dict[int, bytes] = {
        1: b"<< /Type /Catalog /Pages 2 0 R >>",
        3: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /FirstChar 33 "
        b"/LastChar 126 /Widths [%s] /ToUnicode 4 0 R >>" % b" ".join([b"500"] * 94),
        4: b"<< /Length %d >>\nstream\n%s\nendstream" % (len(cmap), cmap),
    }
    page_ids = [6 + 2 * i for i in range(len(pages))]
    objs[2] = b"<< /Type /Pages /Count %d /Kids [%s] >>" % (
        len(pages),
        b" ".join(b"%d 0 R" % p for p in page_ids),
    )
    for i, runs in enumerate(pages):
        stream = b"\n".join(
            b"BT /F1 %d Tf 1 0 0 1 %d %d Tm (%s) Tj ET"
            % (size, x, page_height - top, bytes(codes[c] for c in text))
            for x, top, size, text in runs
        )
        objs[5 + 2 * i] = b"<< /Length %d >>\nstream\n%s\nendstream" % (len(stream), stream)
        objs[6 + 2 * i] = (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 %d] "
            b"/Resources << /Font << /F1 3 0 R >> >> /Contents %d 0 R >>"
            % (page_height, 5 + 2 * i)
        )

    out = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for num in sorted(objs):
        offsets[num] = len(out)
        out += b"%d 0 obj\n" % num + objs[num] + b"\nendobj\n"
    xref_offset, size = len(out), max(objs) + 1
    out += b"xref\n0 %d\n0000000000 65535 f \n" % size
    for num in range(1, size):
        out += b"%010d 00000 n \n" % offsets.get(num, 0)
    out += b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n" % (size, xref_offset)
    path.write_bytes(bytes(out))


def test_text_parser_detects_headings_and_lists(tmp_path):
    path = tmp_path / "doc.md"
    path.write_text("# Title\n\nSome paragraph.\n\n- item one\n- item two\n", encoding="utf-8")

    parsed = get_parser("md").parse(str(path))

    assert parsed.blocks[0].block_type == "heading"
    assert parsed.blocks[0].text == "Title"
    assert any(b.block_type == "list_item" and b.text == "item one" for b in parsed.blocks)
    assert all(b.section == "Title" for b in parsed.blocks[1:])


def test_html_parser_extracts_headings_paragraphs_lists_and_cells(tmp_path):
    path = tmp_path / "doc.html"
    path.write_text(
        "<h1>Intro</h1><p>Hello world</p><ul><li>a point</li></ul>"
        "<table><tr><td>cell</td></tr></table><script>ignored()</script>",
        encoding="utf-8",
    )

    parsed = get_parser("html").parse(str(path))
    types = {b.block_type for b in parsed.blocks}

    assert parsed.blocks[0].block_type == "heading"
    assert {"heading", "paragraph", "list_item", "table_cell"} <= types
    assert not any("ignored" in b.text for b in parsed.blocks)


def test_html_parser_keeps_words_apart_around_inline_markup(tmp_path):
    """get_text(strip=True) concatenates the stripped pieces, so real-world
    markup ('Hello <b>world</b>') comes back as 'Helloworld' - unsearchable."""
    path = tmp_path / "doc.html"
    path.write_text("<p>Hello <b>world</b> and <i>friends</i></p>", encoding="utf-8")

    [block] = get_parser("html").parse(str(path)).blocks

    assert block.text == "Hello world and friends"


def test_html_parser_does_not_emit_nested_tags_twice(tmp_path):
    """<td><p>x</p></td> must yield one block, not the same text as both a
    table_cell and a paragraph - duplicates get indexed and retrieved twice."""
    path = tmp_path / "doc.html"
    path.write_text("<table><tr><td><p>only once</p></td></tr></table>", encoding="utf-8")

    blocks = get_parser("html").parse(str(path)).blocks

    assert [(b.text, b.block_type) for b in blocks] == [("only once", "table_cell")]


def test_a_configured_section_marker_makes_a_line_a_heading():
    """The pattern is the COLLECTION's, not the parser's. Without one these lines
    are body text - too long, Hangul, body font - and the 931 markers of
    유사상품 심사기준 are swallowed into the previous section's paragraph, which is
    exactly where they used to go."""
    import re

    from app.rag.chunking import resolve

    marker = resolve(
        {"strategy": "classification_table", "preset": "korean_ip_classification"}
    ).marker
    line = "[제9류/G390702, G390799] 무선 이어셋, 헤드셋"

    assert _is_heading(line, "", marker) is True
    # No configured pattern means no such rule; the heuristics alone reject it.
    assert _is_heading(line, "") is False
    # Somebody else's pattern does not invent headings in this document either.
    assert _is_heading(line, "", re.compile(r"^SKU-\d+")) is False


def test_pdf_heading_heuristic_accepts_real_headings():
    assert _is_heading("3.2 Results", "Body text follows.") is True
    assert _is_heading("METHODOLOGY", "Body text follows.") is True
    assert _is_heading("Executive Summary", "") is True


def test_pdf_heading_heuristic_accepts_separated_and_multilevel_numbers():
    assert _is_heading("1. Introduction", "Body text follows.") is True
    assert _is_heading("4) Methods", "Body text follows.") is True


def test_pdf_heading_heuristic_rejects_wrapped_body_lines():
    # A wrapped sentence fragment also lacks terminal punctuation - lower-cased
    # words are what keep this from becoming a heading.
    assert _is_heading("the results of the experiment were", "consistent across runs.") is False
    assert _is_heading("This is a complete sentence.", "") is False
    assert _is_heading("x" * 120, "") is False


def test_pdf_heading_heuristic_rejects_numeric_leading_prose():
    """A bare leading number matched before, so a year- or count-leading
    sentence became current_section and relabelled every citation after it."""
    assert _is_heading("2025 was a strong year for the company", "Revenue rose.") is False
    assert _is_heading("15 growers reported blight in June", "Most recovered.") is False
    assert _is_heading("3 of the 12 plots were affected", "The rest were clean.") is False


def test_pdf_parser_emits_headings_with_pages_and_sections(tmp_path):
    """pypdf never emits blank lines between lines of a page, so heading
    detection has to survive on the text alone."""
    path = tmp_path / "doc.pdf"
    _write_pdf(
        path,
        [
            [
                "ANNUAL REPORT",
                "1. Introduction",
                "This document describes the operating results for the fiscal",
                "year and comments on the segments that grew most.",
                "Executive Summary",
                "Revenue grew twelve percent year over year, driven primarily",
                "by the enterprise segment.",
            ],
            ["3.2 Results", "The results were consistent across all runs."],
        ],
    )

    parsed = get_parser("pdf").parse(str(path))
    headings = [b for b in parsed.blocks if b.block_type == "heading"]

    assert [b.text for b in headings] == [
        "ANNUAL REPORT",
        "1. Introduction",
        "Executive Summary",
        "3.2 Results",
    ]
    assert [b.page for b in headings] == [1, 1, 1, 2]
    body = [b for b in parsed.blocks if b.block_type == "paragraph"]
    assert body[0].section == "1. Introduction"
    assert body[-1].section == "3.2 Results"
    assert body[-1].page == 2


def test_pdf_parser_without_headings_yields_one_block_per_page(tmp_path):
    """No structure to find, so the page is the boundary - and the page number
    survives for citations. Task 9's size pass splits these further."""
    path = tmp_path / "plain.pdf"
    _write_pdf(
        path,
        [
            ["Tomato blight spreads through infected soil and splashing water.", "It is bad."],
            ["Growers should rotate crops and remove all infected plant debris.", "So do that."],
        ],
    )

    blocks = get_parser("pdf").parse(str(path)).blocks

    assert [b.block_type for b in blocks] == ["paragraph", "paragraph"]
    assert [b.page for b in blocks] == [1, 2]
    assert blocks[0].text.startswith("Tomato blight")


def test_pdf_heading_heuristic_rejects_uncased_korean_lines():
    """str.isupper() and str.istitle() are Latin case tests and both answer True
    for a line with no cased characters. That is where the shipped chunks'
    garbage section string came from: 691 of 703 isupper() hits over the
    854-page Korean corpus were ordinary body prose or table rows."""
    assert _is_heading("구성C (3) 기재된 위치 C' ( )", "다음 표와 같다") is False
    assert _is_heading("명세서 A B - C", "국어번역문 - B C -") is False
    assert _is_heading("현재 우리나라가 특허제도와 관련하여 가입한 조약은", "설립협약") is False


def test_pdf_heading_heuristic_rejects_contents_entries():
    """A contents line carries the same "3.2 Title" shape as the heading it
    points at; the leader run is what tells them apart."""
    assert _is_heading("4.2 국가 또는 지방자치단체의 권리능력······· 1105", "4.3 법인격이") is False


def test_pdf_parser_rebuilds_lines_in_visual_order(tmp_path):
    """The numerals are drawn AFTER the Hangul run they sit inside, which is how
    Korean government PDFs are produced. Read in content-stream order they land
    at the end of the line ("...연장할 수 있고 교통이 불 1 30 , 편한"); sorting
    the words by (top, x0) puts them back where they belong."""
    path = tmp_path / "scrambled.pdf"
    _write_positioned_pdf(
        path,
        [
            [
                (72, 100, 12, "청구기간은"),
                (134, 100, 12, "회에"),
                (152, 100, 12, "한하여"),
                (128, 100, 12, "1"),  # drawn last, positioned inside the run
            ]
        ],
    )

    [block] = get_parser("pdf").parse(str(path)).blocks

    assert block.text == "청구기간은 1회에 한하여"


def test_pdf_parser_joins_korean_wrapped_words_without_a_space(tmp_path):
    """Korean wraps mid-word with no hyphen, so a space-joined line pair reads
    "교통이 불 편한" and the phrase is unsearchable. The producer emits a
    trailing space glyph when the wrap DID fall on a word boundary, and that
    beats the Hangul rule where it exists - measured, 4,499 of 12,675
    Hangul-to-Hangul breaks in the reference document are real spaces."""
    path = tmp_path / "wrapped.pdf"
    _write_positioned_pdf(
        path,
        [
            [
                (72, 100, 12, "교통이"),
                (114, 100, 12, "불"),
                (72, 130, 12, "편한"),
                (96, 130, 12, "지역에"),
                (72, 160, 12, "보정서를"),
                (108, 160, 12, "제출한"),
                (126, 160, 12, " "),  # trailing space: this wrap IS a word break
                (72, 190, 12, "경우에는"),
                (72, 220, 12, "experiment"),
                (72, 250, 12, "were"),
            ]
        ],
    )

    [block] = get_parser("pdf").parse(str(path)).blocks

    assert "교통이 불편한 지역에" in block.text
    assert "제출한 경우에는" in block.text
    assert "experiment were" in block.text


def test_pdf_parser_keeps_words_apart_across_a_narrow_space_glyph(tmp_path):
    """유사상품 심사기준 sets its table cells with a space narrower than
    WORD_X_TOLERANCE, so a gap-only rule glued every word in them together
    ("불을끄거나예방하기위한기기나장치") - 941 of 1277 sampled joins. The blank
    glyph between the two words is what says a space belongs there. The second
    pair is the control: same sub-tolerance gap, no blank glyph, still glued -
    that is the mixed-font "업(業)으로" case the tolerance was raised for."""
    path = tmp_path / "narrow.pdf"
    # size 4, so each glyph is 2pt wide and the space between "불을" and "끄거나"
    # is a 2pt gap - under WORD_X_TOLERANCE.
    _write_positioned_pdf(
        path,
        [
            [
                (72, 100, 4, "불을"),
                (76, 100, 4, " "),
                (78, 100, 4, "끄거나"),
                (72, 130, 4, "업("),
                (76, 130, 4, "業"),
                (78, 130, 4, ")으로"),
            ]
        ],
    )

    blocks = get_parser("pdf").parse(str(path)).blocks
    text = " ".join(block.text for block in blocks)

    assert "불을 끄거나" in text
    assert "업(業)으로" in text


def _word(x0: float, x1: float, top: float, text: str) -> dict:
    return {"x0": x0, "x1": x1, "top": top, "text": text, "size": 10.0, "fontname": "F"}


def test_columns_split_at_a_gutter_no_glyph_crosses():
    """The two-column half of 유사상품 심사기준's goods table, at its measured
    geometry: a left column ending at x=305, a right column starting at x=321.
    Bucketing by `top` alone read one row of each as a single line - a
    cross-reference from the left column glued to an unrelated goods name from
    the right."""
    words = [
        _word(67, 305, 100, "- 방화피복(제9류/G4507)"),
        _word(321, 560, 100, "가스누출 경보기 gas leak alarms"),
        _word(67, 280, 120, "- 소방차(제12류/G3705)"),
        _word(321, 540, 120, "가스탐지기 gas detecting apparatus"),
    ]

    columns = _columns(words, 624.0)

    assert [[w["text"] for w in column] for column in columns] == [
        ["- 방화피복(제9류/G4507)", "- 소방차(제12류/G3705)"],
        ["가스누출 경보기 gas leak alarms", "가스탐지기 gas detecting apparatus"],
    ]


def test_columns_leaves_a_single_column_page_alone():
    """The control, and the reason the rule is "no glyph on the PAGE crosses it"
    rather than "a wide gap on this line": prose leaves gaps at every ragged line
    end and every indent, and splitting on those would reorder every ordinary
    document. The one full-height band here is the right margin, which is not a
    gutter between columns."""
    words = [
        _word(64, 475, 100, "이것은 한 줄로 이어지는 본문입니다"),
        _word(64, 120, 120, "짧은 줄"),  # ragged: a 355pt gap to its right
        _word(64, 460, 140, "다시 길게 이어지는 본문 한 줄입니다"),
    ]

    assert _columns(words, 539.0) == [words]


def test_pdf_parser_detects_a_heading_by_font_size(tmp_path):
    """Nothing about "overview of results" reads as a heading in plain text -
    lower case, no numbering, no terminal punctuation. Its size does."""
    path = tmp_path / "sized.pdf"
    _write_positioned_pdf(
        path,
        [
            [
                (72, 100, 20, "overview"),
                (162, 100, 20, "of"),
                (192, 100, 20, "results"),
                (72, 150, 12, "revenue"),
                (120, 150, 12, "grew"),
                (156, 150, 12, "twelve"),
                (204, 150, 12, "percent"),
                (72, 180, 12, "across"),
                (120, 180, 12, "every"),
                (156, 180, 12, "reported"),
                (216, 180, 12, "segment."),
            ]
        ],
    )

    blocks = get_parser("pdf").parse(str(path)).blocks

    assert [(b.text, b.block_type) for b in blocks] == [
        ("overview of results", "heading"),
        ("revenue grew twelve percent across every reported segment.", "paragraph"),
    ]
    assert blocks[1].section == "overview of results"


def test_pdf_parser_emits_a_running_header_once_and_drops_folios(tmp_path):
    """A running header is the section name the publisher printed on the page,
    but it is one heading, not one per page - and the folio beside it is not
    text at all. Both were shipped into every chunk of the reference document."""
    path = tmp_path / "running.pdf"
    headers = ["제1부 총 칙"] * 6 + ["제2부 특허출원"] * 6
    _write_positioned_pdf(
        path,
        [
            [
                (72, 50, 12, header),
                (72, 300, 12, f"본문 {i}쪽."),
                (72, 700, 12, str(5300 + i)),
            ]
            for i, header in enumerate(headers)
        ],
    )

    blocks = get_parser("pdf").parse(str(path)).blocks
    headings = [b for b in blocks if b.block_type == "heading"]
    body = [b for b in blocks if b.block_type == "paragraph"]

    assert [(b.text, b.page) for b in headings] == [("제1부 총 칙", 1), ("제2부 특허출원", 7)]
    assert [b.section for b in body] == ["제1부 총 칙"] * 6 + ["제2부 특허출원"] * 6
    assert [b.text for b in body] == [f"본문 {i}쪽." for i in range(12)]


def test_docx_parser_reads_styles_and_tables(tmp_path):
    path = tmp_path / "doc.docx"
    document = DocxDocument()
    document.add_heading("Quarterly Report", level=0)
    document.add_heading("Overview", level=1)
    document.add_paragraph("Revenue grew twelve percent.")
    document.add_paragraph("first bullet", style="List Bullet")
    document.add_paragraph("")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "APAC"
    table.cell(0, 1).text = "120"
    document.save(str(path))

    blocks = get_parser("docx").parse(str(path)).blocks

    assert [(b.text, b.block_type) for b in blocks] == [
        ("Quarterly Report", "heading"),
        ("Overview", "heading"),
        ("Revenue grew twelve percent.", "paragraph"),
        ("first bullet", "list_item"),
        ("APAC", "table_cell"),
        ("120", "table_cell"),
    ]
    assert blocks[2].section == "Overview"


def test_docx_parser_emits_a_merged_cell_once(tmp_path):
    """row.cells repeats the same <w:tc> once per spanned grid column, so a
    merged header cell would be embedded, indexed and retrieved three times."""
    path = tmp_path / "merged.docx"
    document = DocxDocument()
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).merge(table.cell(0, 2)).text = "Regional Summary"
    table.cell(1, 0).text = "APAC"
    table.cell(1, 1).text = "120"
    table.cell(1, 2).text = "up"
    document.save(str(path))

    cells = [b.text for b in get_parser("docx").parse(str(path)).blocks if b.block_type == "table_cell"]

    assert cells == ["Regional Summary", "APAC", "120", "up"]


@pytest.mark.parametrize("file_type,name", [("txt", "a.txt"), ("html", "a.html"), ("pdf", "a.pdf")])
def test_parsers_raise_file_not_found_for_a_missing_file(tmp_path, file_type, name):
    with pytest.raises(FileNotFoundError):
        get_parser(file_type).parse(str(tmp_path / name))


def test_docx_parser_raises_file_not_found_for_a_missing_file(tmp_path):
    """python-docx raises PackageNotFoundError for a missing file, which is not
    a FileNotFoundError - the structure endpoint's 404 branch would miss it."""
    with pytest.raises(FileNotFoundError):
        get_parser("docx").parse(str(tmp_path / "gone.docx"))


def test_get_parser_raises_for_unsupported_type():
    with pytest.raises(ValueError):
        get_parser("exe")


def test_cid_garbage_ratio_flags_unmapped_font_pdfs():
    """(cid:NNNN) 3천 청크가 조용히 색인되던 실사고의 계약: 지배적이면 파싱을
    거절하고, 정상 본문에 낀 소량은 통과시킨다."""
    from app.rag.parsers.pdf_parser import CID_GARBAGE_RATIO, cid_garbage_ratio

    garbage = "(cid:20742)(cid:21106)(cid:22426)" * 100
    assert cid_garbage_ratio(garbage) > CID_GARBAGE_RATIO

    mostly_text = ("정상적인 한국어 본문입니다. " * 200) + "(cid:849)"
    assert cid_garbage_ratio(mostly_text) < CID_GARBAGE_RATIO

    assert cid_garbage_ratio("") == 0.0
